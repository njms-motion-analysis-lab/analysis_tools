from docx import Document

# Create a new Document
doc = Document()

# Add Title and Content to the Document
doc.add_heading('Patient Summary: Latonya, 47-Year-Old Female', level=1)

doc.add_heading('Key Events', level=2)
doc.add_paragraph(
    "• Presentation to ED: SOB, chest + abdominal pain, hypertensive emergency\n"
    "• Contrast Abdominal CT: Steatosis, CHF, cardiomegaly, nonspecific SVC obstruction\n"
    "• ED ECG: Ruled out ACS"
)

doc.add_heading('Chief Complaint', level=2)
doc.add_paragraph(
    "Latonya, a 47-year-old female with a history of ESRD, HTN, HFREF, and SMA syndrome, presented to the ED with SOB, headache, abdominal pain, nausea, and vomiting."
)

doc.add_heading('Subjective', level=2)
doc.add_paragraph(
    "Appearance: Tired and uncomfortable, with widespread edema, but oriented and responsive.\n"
    "Primary Issues: Nausea and vomiting began a day ago, with a headache. Denies medication non-compliance or missed dialysis appointments, assisted by her son at home."
)

doc.add_heading('Objective', level=2)
doc.add_paragraph("Temperature: 36.33°C")

doc.add_heading('Vitals', level=3)
table = doc.add_table(rows=1, cols=7)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Date/Time'
hdr_cells[1].text = 'BP'
hdr_cells[2].text = 'Pulse'
hdr_cells[3].text = 'Resp'
hdr_cells[4].text = 'Temp'
hdr_cells[5].text = 'Weight'
hdr_cells[6].text = 'Height'

# Add Vitals data
vitals_data = [
    ('05/21/24 07:30', '(!) 153/101', '67', '15', '', '', ''),
    ('05/21/24 08:36', '(!) 170/114', '79', '16', '', '', ''),
    ('05/21/24 08:44', '', '74', '', '', '', ''),
    ('05/21/24 11:14', '(!) 190/134', '90', '18', '', '', '')
]

for data in vitals_data:
    row_cells = table.add_row().cells
    for i, item in enumerate(data):
        row_cells[i].text = item

doc.add_heading('CBC', level=3)
table = doc.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Component'
hdr_cells[1].text = 'Value'
hdr_cells[2].text = 'Date'

cbc_data = [
    ('WBC', '3.2 (L)', '05/21/2024'),
    ('RBC', '2.71 (L)', '05/21/2024'),
    ('HGB', '7.7 (L)', '05/21/2024'),
    ('PLT', '264', '05/21/2024'),
    ('MCV', '87.1', '05/21/2024')
]

for data in cbc_data:
    row_cells = table.add_row().cells
    for i, item in enumerate(data):
        row_cells[i].text = item

doc.add_heading('BMP', level=3)
table = doc.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Component'
hdr_cells[1].text = 'Value'
hdr_cells[2].text = 'Date'

bmp_data = [
    ('GLUCOSE', '140 (H)', '05/21/2024'),
    ('SODIUM', '134', '05/21/2024'),
    ('POTASSIUM', '7.0 (HH)', '05/21/2024'),
    ('CHLORIDE', '95 (L)', '05/21/2024'),
    ('CO2', '24', '05/21/2024'),
    ('BUN', '77 (H)', '05/21/2024'),
    ('CREATININE', '8.7 (H)', '05/21/2024'),
    ('CALCIUM', '9.2', '05/21/2024'),
    ('MAGNESIUM', '2.9 (H)', '05/21/2024'),
    ('PHOSPHORUS', '6.3 (H)', '05/21/2024')
]

for data in bmp_data:
    row_cells = table.add_row().cells
    for i, item in enumerate(data):
        row_cells[i].text = item

doc.add_heading('Cardiac Markers', level=3)
table = doc.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Component'
hdr_cells[1].text = 'Value'
hdr_cells[2].text = 'Date'

cardiac_data = [
    ('TROPONIN I', '0.08', '05/20/2024')
]

for data in cardiac_data:
    row_cells = table.add_row().cells
    for i, item in enumerate(data):
        row_cells[i].text = item

doc.add_heading('ABGs', level=3)
table = doc.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Component'
hdr_cells[1].text = 'Value'
hdr_cells[2].text = 'Date'

abgs_data = [
    ('PH', '7.380', '05/20/2024'),
    ('PO2', '61 (H)', '05/20/2024'),
    ('BE', '-1.7', '05/20/2024'),
    ('HCO3', '23 (L)', '05/20/2024')
]

for data in abgs_data:
    row_cells = table.add_row().cells
    for i, item in enumerate(data):
        row_cells[i].text = item

doc.add_heading('LFT', level=3)
table = doc.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Component'
hdr_cells[1].text = 'Value'
hdr_cells[2].text = 'Date'

lft_data = [
    ('TOTALPRO', '6.9', '05/21/2024'),
    ('ASTSGOT', '20', '05/21/2024'),
    ('ALTSGPT', '10', '05/21/2024'),
    ('ALKPHOS', '111 (H)', '05/21/2024'),
    ('ALBUMIN', '3.8', '05/21/2024'),
    ('BILIRUBINTOT', '0.5', '05/21/2024'),
    ('BILIRUBINDIR', '0.1', '05/20/2024')
]

for data in lft_data:
    row_cells = table.add_row().cells
    for i, item in enumerate(data):
        row_cells[i].text = item

doc.add_heading('Assessment', level=2)
doc.add_paragraph(
    "Latonya, a 47-year-old woman with severe hypertension and ESRD, presented to the ED with abdominal pain, nausea, and vomiting, as well as a hypertensive emergency. ECG did not show ACS, but abdominal CT suggested IVC compression. She mentioned a productive cough. She received 50 mg IV nitro, albuterol, toradol, morphine, and zofran. On exam, she was in moderate distress, edematous, and described her main problems as diffuse abdominal pain and nausea. This is around her 10th admission in 2024, with many overlapping issues. Ensuring she can regularly take her prescribed medications is crucial."
)

doc.add_heading('Plan', level=2)
doc.add_paragraph(
    "1. Hypertensive Emergency\n"
    "   a. Presented with BP of 230/176 and evidence of end organ damage (pulmonary edema on CXR). BP improving, currently 140/100, briefly 190 during dialysis.\n"
    "   b. Continue IV nitro until BP stabilizes, then resume home BP medications.\n\n"
    "2. End Stage Renal Disease and Electrolyte Abnormalities\n"
    "   a. Evidence of extensive kidney damage: creatinine high of 8.2, BUN of 77. Currently on dialysis without appointment issues, assisted by her brother.\n"
    "   b. Severe hyperkalemia (7.0) and hyperphosphatemia (6.3).\n"
    "   c. Improved BP control to help manage ESRD. Continue hemodialysis. Follow up with SW for extensive history and potential in-home care.\n\n"
    "3. HFREF\n"
    "   a. CXR showed cardiomegaly and heart failure, emergent imaging indicated SVC compression.\n"
    "   b. Follow up on possible SVC compression with chest CT. History of left arm swelling raises concerns for SVC syndrome.\n"
    "   c. Ensure adherence to home medications for HTN/HF management.\n\n")
