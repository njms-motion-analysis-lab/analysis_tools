import argparse
import nltk
from nltk.tokenize import word_tokenize, sent_tokenize
from nltk.corpus import stopwords
from nltk.stem import PorterStemmer, WordNetLemmatizer
from nltk import pos_tag, ne_chunk
from textblob import TextBlob
from wordcloud import WordCloud
import matplotlib.pyplot as plt
from docx import Document

# Ensure the necessary NLTK data is downloaded
nltk.download('punkt')
nltk.download('stopwords')
nltk.download('averaged_perceptron_tagger')
nltk.download('maxent_ne_chunker')
nltk.download('words')

# Default text if no input is provided
DEFAULT_TEXT = "default_text.docx"

def read_default_text():
    try:
        doc = Document(DEFAULT_TEXT)
        full_text = []
        for paragraph in doc.paragraphs:
            full_text.append(paragraph.text)
        return ' '.join(full_text)
    except Exception as e:
        print(f"Error reading default text: {e}")
        return ""

# Functions for text processing
def tokenize(text):
    words = word_tokenize(text)
    sentences = sent_tokenize(text)
    print("Words:", words)
    print("Sentences:", sentences)

def remove_stopwords(text):
    words = word_tokenize(text)
    stop_words = set(stopwords.words('english'))
    filtered_words = [word for word in words if word.lower() not in stop_words]
    print("Filtered Words:", filtered_words)

def stem(text):
    words = word_tokenize(text)
    stemmer = PorterStemmer()
    stemmed_words = [stemmer.stem(word) for word in words]
    print("Stemmed Words:", stemmed_words)

def lemmatize(text):
    words = word_tokenize(text)
    lemmatizer = WordNetLemmatizer()
    lemmatized_words = [lemmatizer.lemmatize(word) for word in words]
    print("Lemmatized Words:", lemmatized_words)

def pos_tagging(text):
    words = word_tokenize(text)
    pos_tags = pos_tag(words)
    print("POS Tags:", pos_tags)

def named_entity_recognition(text):
    words = word_tokenize(text)
    pos_tags = pos_tag(words)
    named_entities = ne_chunk(pos_tags)
    print("Named Entities:", named_entities)

def sentiment_analysis(text):
    blob = TextBlob(text)
    sentiment = blob.sentiment
    print("Sentiment:", sentiment)

def word_cloud(text):
    wordcloud = WordCloud(width=800, height=400, background_color='white').generate(text)
    plt.figure(figsize=(10, 5))
    plt.imshow(wordcloud, interpolation='bilinear')
    plt.axis('off')
    plt.show()

# Main function to parse arguments and call the appropriate function
def main():
    parser = argparse.ArgumentParser(description="Word Analyzer Tool")
    parser.add_argument("function", choices=["tokenize", "remove_stopwords", "stem", "lemmatize", "pos_tagging", "ner", "sentiment", "wordcloud"], help="Function to execute")
    parser.add_argument("--text", type=str, help="Text to analyze or path to a text file", default=None)
    args = parser.parse_args()

    # Read text from file if path is provided
    if args.text and args.text.endswith(".docx"):
        try:
            doc = Document(args.text)
            full_text = []
            for paragraph in doc.paragraphs:
                full_text.append(paragraph.text)
            text = ' '.join(full_text)
        except Exception as e:
            print(f"Error reading text file: {e}")
            text = read_default_text()
    elif args.text:
        text = args.text
    else:
        text = read_default_text()

    if args.function == "tokenize":
        tokenize(text)
    elif args.function == "remove_stopwords":
        remove_stopwords(text)
    elif args.function == "stem":
        stem(text)
    elif args.function == "lemmatize":
        lemmatize(text)
    elif args.function == "pos_tagging":
        pos_tagging(text)
    elif args.function == "ner":
        named_entity_recognition(text)
    elif args.function == "sentiment":
        sentiment_analysis(text)
    elif args.function == "wordcloud":
        word_cloud(text)

if __name__ == "__main__":
    main()