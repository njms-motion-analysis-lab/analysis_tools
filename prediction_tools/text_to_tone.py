import argparse
import nltk
from nltk.tokenize import word_tokenize, sent_tokenize
from nltk.corpus import stopwords
from nltk.stem import PorterStemmer, WordNetLemmatizer
from nltk import pos_tag, ne_chunk
from textblob import TextBlob
from wordcloud import WordCloud
import matplotlib.pyplot as plt
import seaborn as sns
from collections import Counter
from docx import Document

# Ensure the necessary NLTK data is downloaded
nltk.download('punkt')
nltk.download('stopwords')
nltk.download('averaged_perceptron_tagger')
nltk.download('maxent_ne_chunker')
nltk.download('words')

# Default text if no input is provided
DEFAULT_TEXT = "/Users/stephenmacneille/Desktop/20240616_Trial1.docx"

# Custom stop words
CUSTOM_STOP_WORDS = {
    'yeah', 'oh', 'really', 'okay', 'know', 'gotcha', 'like', 'kind', 'right', 'feel', 'think', 'honestly', 
    'anything', 'want', "n't", "'s", "'m", "'ve", "'ll", "'d", 'ca', 'wo', 're', 'nt', 'na', 'gon'
}

def read_default_text():
    try:
        doc = Document(DEFAULT_TEXT)
        full_text = []
        for paragraph in doc.paragraphs:
            full_text.append(paragraph.text)
        return ' '.join(full_text)
    except Exception as e:
        print(f"Error reading default text: {e}")
        return ""

def remove_stopwords(text):
    words = word_tokenize(text)
    stop_words = set(stopwords.words('english')).union(CUSTOM_STOP_WORDS)
    filtered_words = [word for word in words if word.lower() not in stop_words]
    return ' '.join(filtered_words)

def clean_tokens(tokens):
    # Remove punctuation and non-alphabetic tokens
    tokens = [word for word in tokens if word.isalpha()]
    return tokens

# Functions for text processing
def tokenize(text):
    words = word_tokenize(text)
    words = clean_tokens(words)
    sentences = sent_tokenize(text)
    print("Words:", words)
    print("Sentences:", sentences)

    # Visualization
    word_freq = Counter(words)
    common_words = word_freq.most_common(10)
    words, counts = zip(*common_words)
    plt.figure(figsize=(10, 5))
    sns.barplot(x=list(counts), y=list(words))
    plt.title("Top 10 Most Common Words")
    plt.show()

def top_words(text):
    filtered_text = remove_stopwords(text)
    return tokenize(filtered_text)

def stem(text):
    words = word_tokenize(text)
    words = clean_tokens(words)
    stemmer = PorterStemmer()
    stemmed_words = [stemmer.stem(word) for word in words]
    print("Stemmed Words:", stemmed_words)

def lemmatize(text):
    words = word_tokenize(text)
    words = clean_tokens(words)
    lemmatizer = WordNetLemmatizer()
    lemmatized_words = [lemmatizer.lemmatize(word) for word in words]
    print("Lemmatized Words:", lemmatized_words)

def pos_tagging(text):
    words = word_tokenize(text)
    words = clean_tokens(words)
    pos_tags = pos_tag(words)
    print("POS Tags:", pos_tags)

    # Visualization
    pos_counts = Counter(tag for word, tag in pos_tags)
    tags, counts = zip(*pos_counts.items())
    plt.figure(figsize=(10, 5))
    sns.barplot(x=list(counts), y=list(tags))
    plt.title("Part of Speech Tagging")
    plt.show()

def named_entity_recognition(text):
    words = word_tokenize(text)
    words = clean_tokens(words)
    pos_tags = pos_tag(words)
    named_entities = ne_chunk(pos_tags)
    print("Named Entities:", named_entities)

def sentiment_analysis(text):
    filtered_text = remove_stopwords(text)
    blob = TextBlob(filtered_text)
    sentiment = blob.sentiment
    print("Sentiment:", sentiment)

    # Visualization
    plt.figure(figsize=(10, 5))
    plt.bar(['Polarity', 'Subjectivity'], [sentiment.polarity, sentiment.subjectivity], color=['blue', 'orange'])
    plt.ylim([-1, 1])
    plt.title('Sentiment Analysis')
    plt.show()

def word_cloud(text):
    filtered_text = remove_stopwords(text)
    words = word_tokenize(filtered_text)
    stemmer = PorterStemmer()
    stemmed_words = [stemmer.stem(word) for word in words]
    wordcloud = WordCloud(width=800, height=400, background_color='white').generate(' '.join(stemmed_words))
    plt.figure(figsize=(10, 5))
    plt.imshow(wordcloud, interpolation='bilinear')
    plt.axis('off')
    plt.show()

# Main function to parse arguments and call the appropriate function
def main():
    parser = argparse.ArgumentParser(description="Word Analyzer Tool")
    parser.add_argument("function", choices=["tokenize", "stem", "lemmatize", "pos_tagging", "ner", "sentiment", "wordcloud", "top"], help="Function to execute")
    parser.add_argument("--text", type=str, help="Text to analyze or path to a text file", default=None)
    args = parser.parse_args()

    # Read text from file if path is provided
    if args.text and args.text.endswith(".docx"):
        try:
            doc = Document(args.text)
            full_text = []
            for paragraph in doc.paragraphs:
                full_text.append(paragraph.text)
            text = ' '.join(full_text)
        except Exception as e:
            print(f"Error reading text file: {e}")
            text = read_default_text()
    elif args.text:
        text = args.text
    else:
        text = read_default_text()

    if args.function == "tokenize":
        tokenize(text)
    elif args.function == "stem":
        stem(text)
    elif args.function == "lemmatize":
        lemmatize(text)
    elif args.function == "pos_tagging":
        pos_tagging(text)
    elif args.function == "ner":
        named_entity_recognition(text)
    elif args.function == "sentiment":
        sentiment_analysis(text)
    elif args.function == "wordcloud":
        word_cloud(text)
    elif args.function == "top":
        top_words(text)

if __name__ == "__main__":
    main()